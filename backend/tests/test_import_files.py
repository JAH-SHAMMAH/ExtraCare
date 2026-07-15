"""Multi-format bulk import — the shared parser + the CBT / Library / wizard entry points.

Proves a CSV, an Excel sheet, and a Word table all yield the same row dicts, that a
Word/PDF without a table (and unsupported types) fail with a clear message, and that
the three importers (CBT question bank, Library books, the generic parse-file
endpoint) accept the richer formats — not just CSV.
"""
import io
import uuid

import pytest
import openpyxl
import docx
from fastapi import HTTPException
from starlette.datastructures import UploadFile
from sqlalchemy import select

from app.models.user import User, UserStatus
from app.models.role import Role
from app.models.modules.school import LibraryBook, QuestionBankItem
from app.services.import_files import rows_from_upload
from app.routers.modules.cbt import import_bank
from app.routers.modules.library import import_books
from app.routers.imports import parse_file

pytestmark = pytest.mark.asyncio


def _xlsx(headers, *rows) -> bytes:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(list(headers))
    for r in rows:
        ws.append(list(r))
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _docx_table(headers, *rows) -> bytes:
    d = docx.Document()
    t = d.add_table(rows=1 + len(rows), cols=len(headers))
    for j, h in enumerate(headers):
        t.cell(0, j).text = str(h)
    for i, r in enumerate(rows, start=1):
        for j, v in enumerate(r):
            t.cell(i, j).text = str(v)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _pdf_table(headers, *rows) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Table
    from reportlab.lib import colors
    buf = io.BytesIO()
    t = Table([list(headers), *[list(r) for r in rows]])
    t.setStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.black)])
    SimpleDocTemplate(buf, pagesize=letter).build([t])
    return buf.getvalue()


def _docx_notable() -> bytes:
    d = docx.Document()
    d.add_paragraph("Just prose, no table.")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _upload(name, content) -> UploadFile:
    return UploadFile(io.BytesIO(content), filename=name)


async def _admin(db, org) -> User:
    u = User(id=str(uuid.uuid4()), email=f"a-{uuid.uuid4().hex[:6]}@x.com", full_name="Admin",
             status=UserStatus.ACTIVE, org_id=org.id)
    role = Role(id=str(uuid.uuid4()), name="org_admin", slug="org_admin", permissions=[], org_id=org.id, is_system=False)
    db.add(role)
    u.roles = [role]
    db.add(u)
    await db.commit()
    return u


# ── The shared parser ──────────────────────────────────────────────────────────

async def test_csv_xlsx_docx_yield_same_rows():
    headers = ("title", "author")
    expected = [{"title": "Dune", "author": "Herbert"}, {"title": "IT", "author": "King"}]
    csv_bytes = b"title,author\nDune,Herbert\nIT,King\n"
    assert rows_from_upload("x.csv", csv_bytes) == expected
    assert rows_from_upload("x.xlsx", _xlsx(headers, ("Dune", "Herbert"), ("IT", "King"))) == expected
    assert rows_from_upload("x.docx", _docx_table(headers, ("Dune", "Herbert"), ("IT", "King"))) == expected
    assert rows_from_upload("x.pdf", _pdf_table(headers, ("Dune", "Herbert"), ("IT", "King"))) == expected


async def test_parser_error_paths():
    with pytest.raises(ValueError):
        rows_from_upload("notes.docx", _docx_notable())        # no table
    with pytest.raises(ValueError):
        rows_from_upload("photo.png", b"\x89PNG...")            # unsupported type
    with pytest.raises(ValueError):
        rows_from_upload("empty.csv", b"")                     # empty file


# ── CBT question-bank import (Excel) ────────────────────────────────────────────

async def test_cbt_bank_import_accepts_xlsx(db, org):
    admin = await _admin(db, org)
    content = _xlsx(
        ("question", "type", "option_a", "option_b", "correct_answer", "points"),
        ("2 + 2 = ?", "mcq", "3", "4", "B", "1"),
        ("Capital of France?", "mcq", "Paris", "Rome", "A", "2"),
    )
    res = await import_bank(file=_upload("bank.xlsx", content), request=None, db=db, current_user=admin)
    assert res["imported"] == 2
    n = (await db.execute(select(QuestionBankItem).where(QuestionBankItem.org_id == org.id))).scalars().all()
    assert len(n) == 2


# ── Library book import (Word table) ────────────────────────────────────────────

async def test_library_book_import_accepts_docx(db, org):
    admin = await _admin(db, org)
    content = _docx_table(
        ("title", "author", "total_copies", "category"),
        ("Things Fall Apart", "Achebe", "3", "Fiction"),
        ("", "NoTitle", "1", "Fiction"),          # missing title → skipped
    )
    res = await import_books(file=_upload("books.docx", content), db=db, current_user=admin)
    assert res["imported"] == 1 and len(res["errors"]) == 1
    book = (await db.execute(select(LibraryBook).where(LibraryBook.org_id == org.id))).scalar_one()
    assert book.title == "Things Fall Apart" and book.total_copies == 3 and book.available_copies == 3


# ── Generic wizard parse endpoint ───────────────────────────────────────────────

async def test_parse_file_endpoint_returns_headers_and_rows(db, org):
    admin = await _admin(db, org)
    content = _xlsx(("first_name", "last_name"), ("Ada", "Okafor"), ("Zed", "X"))
    res = await parse_file(file=_upload("students.xlsx", content), current_user=admin)
    assert res["headers"] == ["first_name", "last_name"]
    assert res["rows"] == [{"first_name": "Ada", "last_name": "Okafor"}, {"first_name": "Zed", "last_name": "X"}]

    with pytest.raises(HTTPException) as ei:
        await parse_file(file=_upload("notes.docx", _docx_notable()), current_user=admin)
    assert ei.value.status_code == 422
