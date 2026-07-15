"""Multi-format tabular import — turn an uploaded CSV / Excel / Word / PDF into a
list of ``{header: value}`` row dicts.

CSV + XLSX read the sheet directly. DOCX + PDF read the tables in the document
(the first row is the header); a Word/PDF WITHOUT a table can't be parsed into
records and raises a clear message. Callers keep their own field mapping — this
only handles the file → rows step, so every importer supports every format from
one place without duplicating parsing.
"""
from __future__ import annotations

import csv
import io

SUPPORTED_EXTS = (".csv", ".xlsx", ".docx", ".pdf")
# For the frontend file picker + humans.
ACCEPT_ATTR = ".csv,.xlsx,.docx,.pdf"


def _ext(filename: str) -> str:
    name = (filename or "").lower().strip()
    for e in SUPPORTED_EXTS:
        if name.endswith(e):
            return e
    return ""


def _rows_from_grid(grid: list[list[str]]) -> list[dict]:
    """First non-empty row = headers; each later row = a record keyed by header.
    Blank rows are dropped; short rows are padded with empty strings."""
    rows = [r for r in grid if any((c or "").strip() for c in r)]
    if not rows:
        return []
    headers = [(c or "").strip() for c in rows[0]]
    out: list[dict] = []
    for r in rows[1:]:
        rec: dict[str, str] = {}
        for i, h in enumerate(headers):
            if h:
                rec[h] = (str(r[i]).strip() if i < len(r) and r[i] is not None else "")
        if rec:
            out.append(rec)
    return out


def _csv_rows(content: bytes) -> list[dict]:
    try:
        text = content.decode("utf-8-sig")
    except UnicodeDecodeError:
        raise ValueError("CSV must be UTF-8 encoded.")
    reader = csv.reader(io.StringIO(text))
    return _rows_from_grid([list(r) for r in reader])


def _xlsx_rows(content: bytes) -> list[dict]:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    try:
        ws = wb.active
        grid = [["" if c is None else str(c) for c in row] for row in ws.iter_rows(values_only=True)]
    finally:
        wb.close()
    return _rows_from_grid(grid)


def _docx_rows(content: bytes) -> list[dict]:
    import docx

    doc = docx.Document(io.BytesIO(content))
    if not doc.tables:
        raise ValueError("The Word document has no table. Put the data in a table whose first row is the column headers.")
    grid = [[cell.text for cell in row.cells] for row in doc.tables[0].rows]
    return _rows_from_grid(grid)


def _pdf_rows(content: bytes) -> list[dict]:
    import pdfplumber

    grid: list[list[str]] = []
    header: list[str] | None = None
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            for table in (page.extract_tables() or []):
                for row in table:
                    cells = ["" if c is None else str(c).strip() for c in row]
                    if header is None:
                        header = cells
                        grid.append(cells)
                    elif cells == header:
                        continue  # a repeated header row on a later page of the same table
                    else:
                        grid.append(cells)
    if not grid:
        raise ValueError("No table found in the PDF. Provide a PDF containing a table whose first row is the column headers.")
    return _rows_from_grid(grid)


def rows_from_upload(filename: str, content: bytes) -> list[dict]:
    """Extract row dicts from an uploaded file, picking the parser by extension.
    Raises ``ValueError`` (message safe to surface to the user) on an unsupported
    type, an empty file, or a Word/PDF with no table."""
    ext = _ext(filename)
    if not ext:
        raise ValueError("Unsupported file type. Upload a CSV, Excel (.xlsx), Word (.docx) or PDF file.")
    if not content:
        raise ValueError("The file is empty.")
    if ext == ".csv":
        return _csv_rows(content)
    if ext == ".xlsx":
        return _xlsx_rows(content)
    if ext == ".docx":
        return _docx_rows(content)
    return _pdf_rows(content)
